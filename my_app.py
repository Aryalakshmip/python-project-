from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture('pexels-lola-russian-1855582.jpg',width=Inches(1.0))

# name phone number and email details
name= input("What is your name ?")
speak('hello'+name+'how are you today')
phone_number=input("What is your phone number ?")
email=input("What is your email ?")

document.add_paragraph(name + ' | '+ phone_number + ' | '+ email)
# about me

document.add_heading('About Me')
about_me = input('Tell about yourself ?')
document.add_paragraph(about_me)

# skills
document.add_heading('Skills')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style ='List Bullet'

while True:
    has_more_skills = input('do you have more skills? Yes or No')
    if has_more_skills.lower() == 'yes':
        skill  =input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break    


# work experience

document.add_heading("Work Experience")
p = document.add_paragraph()

company = input('enter company?')
from_date = input('from date')
to_date = input('To Date')

p.add_run(company+ ' ' ).bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input( 'Describe your experience at ' + company)
p.add_run(experience_details)

# more experiences
while True :
    has_more_experiences = input('do you have more experience? Yes or No')
    if has_more_experiences.lower() =='yes':
        p = document.add_paragraph()

        company = input('enter company?')
        from_date = input('from date')
        to_date = input('To Date')

        p.add_run(company+ ' ' ).bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input( 'Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break

document.save('cv.docx')